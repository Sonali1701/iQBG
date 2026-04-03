import os
import shutil
import zipfile
from pathlib import Path
from typing import Dict, List

from PIL import Image
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas

from xlsx_utils import build_tagging_xlsx


def _set_two_columns(section, num_cols: int = 2, space_twips: int = 720) -> None:
    sect_pr = section._sectPr
    cols_nodes = sect_pr.xpath("./w:cols")
    cols = cols_nodes[0] if cols_nodes else OxmlElement("w:cols")
    cols.set(qn("w:num"), str(num_cols))
    cols.set(qn("w:space"), str(space_twips))
    if not cols_nodes:
        sect_pr.append(cols)


def build_word_two_column(image_paths: List[Path], out_path: Path, title: str) -> None:
    doc = Document()
    section = doc.sections[0]
    _set_two_columns(section)

    doc.add_heading(title, level=1)
    max_width = Inches(3.1)

    for image_path in image_paths:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(str(image_path), width=max_width)

    doc.save(out_path)


def build_pdf_two_column(image_paths: List[Path], out_path: Path, title: str) -> None:
    page_width, page_height = A4
    margin = 36
    gutter = 18
    title_gap = 28
    col_width = (page_width - (2 * margin) - gutter) / 2
    x_positions = [margin, margin + col_width + gutter]
    usable_height = page_height - margin - title_gap

    pdf = canvas.Canvas(str(out_path), pagesize=A4)
    pdf.setTitle(title)
    pdf.setFont("Helvetica-Bold", 14)
    pdf.drawString(margin, page_height - margin + 4, title)

    col_idx = 0
    y_top = page_height - margin - title_gap
    y = y_top

    for image_path in image_paths:
        reader = ImageReader(str(image_path))
        with Image.open(image_path) as image:
            width_px, height_px = image.size

        if not width_px or not height_px:
            continue

        scale = min(col_width / width_px, usable_height / height_px)
        draw_width = width_px * scale
        draw_height = height_px * scale

        if y - draw_height < margin:
            if col_idx == 0:
                col_idx = 1
                y = y_top
            else:
                pdf.showPage()
                pdf.setFont("Helvetica-Bold", 14)
                pdf.drawString(margin, page_height - margin + 4, title)
                col_idx = 0
                y = y_top

        x = x_positions[col_idx] + (col_width - draw_width) / 2
        y_draw = y - draw_height
        pdf.drawImage(reader, x, y_draw, width=draw_width, height=draw_height, preserveAspectRatio=True, mask="auto")
        y = y_draw - 16

    pdf.save()


def _build_xlsx_rows(selected_questions: List[Dict[str, str]]) -> List[List[str]]:
    rows: List[List[str]] = []
    for order, item in enumerate(selected_questions, start=1):
        row = [""] * 25
        row[0] = str(order)
        row[1] = item.get("subject", "")
        row[2] = item.get("chapter", "")
        row[3] = item.get("topic", "")
        row[4] = item.get("subtopic", "")
        row[6] = item["export_filename"]
        row[13] = item.get("question_type", "Single Choice")
        row[14] = "4"
        row[15] = "1"
        row[17] = item.get("difficulty", "Medium")
        row[18] = "English"
        row[22] = f"Q{order}"
        rows.append(row)
    return rows


def prepare_export_assets(job_dir: Path, selected_questions: List[Dict[str, str]], project_name: str) -> Dict[str, str]:
    try:
        from ppt_builder import generate_ppts
    except ModuleNotFoundError as exc:
        raise RuntimeError("PPT export dependency is missing. Install python-pptx before finalizing exports.") from exc

    export_dir = job_dir / "final_output"
    images_dir = export_dir / "images"
    images_dir.mkdir(parents=True, exist_ok=True)

    export_items: List[Dict[str, str]] = []
    image_paths: List[Path] = []

    for order, item in enumerate(selected_questions, start=1):
        target_name = f"QUES_ENG_Q{order}.png"
        target_path = images_dir / target_name
        shutil.copy2(item["image_path"], target_path)

        export_item = dict(item)
        export_item["export_filename"] = target_name
        export_item["export_path"] = str(target_path)
        export_items.append(export_item)
        image_paths.append(target_path)

    rows = _build_xlsx_rows(export_items)
    xlsx_path = export_dir / f"{project_name or 'Paper'}_Tagging.xlsx"
    build_tagging_xlsx(rows, xlsx_path, "English", include_tagging=True)

    ppt_files = generate_ppts(str(export_dir), "english", True)
    word_path = export_dir / f"{project_name or 'Paper'}_Paper.docx"
    pdf_path = export_dir / f"{project_name or 'Paper'}_Paper.pdf"

    build_word_two_column(image_paths, word_path, project_name or "Generated Paper")
    build_pdf_two_column(image_paths, pdf_path, project_name or "Generated Paper")

    zip_path = job_dir / f"{project_name or 'paper'}_bundle.zip"
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as bundle:
        for root, _, files in os.walk(export_dir):
            for file_name in files:
                full_path = Path(root) / file_name
                bundle.write(full_path, full_path.relative_to(export_dir))

    return {
        "export_dir": str(export_dir),
        "xlsx_path": str(xlsx_path),
        "word_path": str(word_path),
        "pdf_path": str(pdf_path),
        "zip_path": str(zip_path),
        "ppt_paths": ppt_files,
    }
